import os
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter
from docx import Document
from convertor import convertValue
from datetime import datetime
import logging
workbook = Workbook()
sheet = workbook.active
table_sheet = workbook.create_sheet("Tables")
allDocuments = []
headers = ["File"]
row = 2
row_number = 1
time = datetime.now()
file_name = f"-{time.hour}-{time.minute}-{time.second}"
logging.basicConfig(
    filename=f"logs{file_name}.log",              
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
def checkParagraphs(document):
    data = {}
    data["File"] = (file,None)
    for paragraph in document.paragraphs:
        text = paragraph.text
        try:
            key, value = text.split(":", 1)
            key = key.strip()
            value = value.strip()
            convValue, method = convertValue(value)
            data[key] = (convValue, method)
            if key not in headers:
                headers.append(key)
        except ValueError as error:
            logging.error(f"Error in {file}: {error} | Text: {text!r}")
            continue
    return data
def checkTables(document):
    tables = []
    for table in document.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                value, method = convertValue(cell.text)
                row_data.append((value,method))
            table_data.append(row_data)
        tables.append(table_data)
    return tables
def writeTables(tables, sheet, row_number):
    for table in tables:
        for row in table:
            for column, (value,method) in enumerate(row, start=1):
                cell = sheet.cell(row=row_number, column=column)
                cell.value = value
                if method == "money$":
                    cell.number_format = "$#,##0.00"
                if method == "money€":
                    cell.number_format = "#,##0.00€"
                if method == "dateDDMMYYYY":
                    cell.number_format = "DD.MM.YYYY"
            row_number += 1
    return row_number
for file in os.listdir("OfficeFiles"):
    if file.endswith(".docx"):
        filepath = os.path.join("OfficeFiles", file)
        try:
            document = Document(filepath)
        except Exception as error:
            logging.error(f"Can't open {file}: {error}")
            continue
        data = checkParagraphs(document)
        allDocuments.append(data)
        row_number = writeTables(checkTables(document), table_sheet, row_number)
        logging.info(f"{file} Done")
for column, header in enumerate(headers, start=1):
        excelCell = sheet.cell(row=1, column=column)
        excelCell.value = header
        excelCell.font = Font(bold=True)
for data in allDocuments:   
    for column, header in enumerate(headers, start=1):
        result = data.get(header)
        if result is None:
            continue    
        value,method = result
        excelCell = sheet.cell(row=row, column=column)
        excelCell.value = value
        if method == "money$":
            excelCell.number_format = "$#,##0.00"
        if method == "money€":
            excelCell.number_format = "#,##0.00€"
        if method == "dateDDMMYYYY":
            excelCell.number_format = "DD.MM.YYYY"
    row += 1
for column in sheet.columns:
    max_length = 0
    for cell in column:
        if cell.value is not None:
            max_length = max(max_length, len(str(cell.value)))
    column_letter = get_column_letter(column[0].column)
    sheet.column_dimensions[column_letter].width = max_length + 5
sheet.auto_filter.ref = sheet.dimensions
sheet.freeze_panes = "A2"
for column in table_sheet.columns:
    max_length = 0
    for cell in column:
        if cell.value is not None:
            max_length = max(max_length, len(str(cell.value)))
    column_letter = get_column_letter(column[0].column)
    table_sheet.column_dimensions[column_letter].width = max_length + 5
table_sheet.auto_filter.ref = table_sheet.dimensions
table_sheet.freeze_panes = "A2"
workbook.save(f"OfficeFiles/Result{file_name}.xlsx")